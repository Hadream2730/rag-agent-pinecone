import os
import fitz                                  
from docx import Document as DocxDocument    
from typing import List, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain.docstore.document import Document
from pinecone import Pinecone, ServerlessSpec
import concurrent.futures, itertools

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
EMBEDDING_MODEL = "text-embedding-3-small"
CHUNK_SIZE      = 1500
CHUNK_OVERLAP   = 80
MAX_INPUT_TOKENS = 300000

def load_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        pdf = fitz.open(path)
        return "\n".join(page.get_text() for page in pdf)
    elif ext == ".docx":
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def create_pinecone_index(
    paths: List[str],
    index_name: str | None = None,
    progress_cb = None,
) -> Any:
   
    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    all_chunks = []
    for p in paths:
        raw = load_text(p)
        docs = [Document(page_content=raw, metadata={"source": os.path.basename(p)})]
        chunks = splitter.split_documents(docs)
        all_chunks.extend(chunks)
    total_chunks = len(all_chunks)
    print(f"[EmbeddingCreator] total chunks: {total_chunks}")
    if progress_cb:
        progress_cb(5)  # initial step after loading and splitting

    batch_size = MAX_INPUT_TOKENS // CHUNK_SIZE
    embedder = OpenAIEmbeddings(model=EMBEDDING_MODEL, api_key=OPENAI_API_KEY)

    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))

    index_name = os.getenv("PINECONE_INDEX_NAME", "rag-agent-index")

    # Wipe previous index 
    if pc.has_index(index_name):
        pc.delete_index(index_name)

    # Serverless spec parameters
    cloud = os.getenv("PINECONE_CLOUD", "aws")
    region = os.getenv("PINECONE_REGION", "us-east-1")

    pc.create_index(
        name=index_name,
        dimension=1536,  
        metric="cosine",
        spec=ServerlessSpec(cloud=cloud, region=region),
    )

    index = pc.Index(index_name)

    # Prepare batches
    batches = [all_chunks[i : i + batch_size] for i in range(0, len(all_chunks), batch_size)]

    def process_batch(batch_tuple):
        idx_offset, batch = batch_tuple
        texts = [c.page_content for c in batch]
        metas = [c.metadata for c in batch]
        embeddings = embedder.embed_documents(texts)
        records = [
            {
                "id": f"doc-{idx_offset + j}",
                "values": embeddings[j],
                "metadata": {**metas[j], "text": texts[j]},
            }
            for j in range(len(texts))
        ]
        index.upsert(records)
        return len(batch)

    processed = 0
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        for count in executor.map(process_batch, [(i*batch_size, b) for i, b in enumerate(batches)]):
            processed += count
            if progress_cb:
                pct = 5 + int(90 * processed / total_chunks)
                progress_cb(min(pct, 95))
            print(f"[EmbeddingCreator] indexed {count} chunks concurrently")

    if progress_cb:
        progress_cb(100)
    print(f"[EmbeddingCreator] Pinecone index '{index_name}' populated and ready.")
    return index
